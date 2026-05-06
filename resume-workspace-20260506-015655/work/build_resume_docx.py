from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WORKSPACE = Path(__file__).resolve().parents[1]
OUT = WORKSPACE / "output" / "Carlton_Shi_AI_Creative_Content_Strategy_Resume.docx"


ACCENT = RGBColor(31, 78, 121)
MUTED = RGBColor(90, 90, 90)
BODY = RGBColor(30, 30, 30)


def set_cell_margin(section):
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)


def set_paragraph_border(paragraph, color="D9E2F3", size="6"):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_run(paragraph, text, bold=False, italic=False, size=None, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    return run


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    add_run(p, text.upper(), bold=True, size=9.5, color=ACCENT)
    set_paragraph_border(p)
    return p


def add_role(doc, org, location, role, dates):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    add_run(p, org, bold=True, size=9.5, color=BODY)
    add_run(p, f" - {location}", size=9.5, color=BODY)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(1)
    p2.paragraph_format.keep_with_next = True
    add_run(p2, role, italic=True, size=9, color=MUTED)
    add_run(p2, f" | {dates}", italic=True, size=9, color=MUTED)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1.6)
    p.paragraph_format.line_spacing = 1.02
    for part in text:
        if isinstance(part, tuple):
            add_run(p, part[0], bold=part[1], size=8.7, color=BODY)
        else:
            add_run(p, part, size=8.7, color=BODY)
    return p


def add_skill_line(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1.7)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, label + ": ", bold=True, size=8.8, color=BODY)
    add_run(p, text, size=8.8, color=BODY)


def build():
    doc = Document()
    set_cell_margin(doc.sections[0])

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(8.8)
    styles["Normal"].font.color.rgb = BODY

    if "List Bullet" in styles:
        styles["List Bullet"].font.name = "Arial"
        styles["List Bullet"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        styles["List Bullet"].font.size = Pt(8.7)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(1)
    add_run(title, "Jingchuan (Carlton) Shi", bold=True, size=18, color=ACCENT)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(1)
    add_run(
        subtitle,
        "AI Creative & Creative Technologist | Generative Media | Film Marketing | Content Strategy",
        size=9.2,
        color=BODY,
    )

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(5)
    add_run(contact, "315-395-2313 | jshi77@syr.edu | LinkedIn | Portfolio", size=8.8, color=MUTED)

    add_section_heading(doc, "Profile")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.03
    add_run(
        p,
        "Creative technologist and filmmaker working across generative AI, VFX, interactive web, and film marketing. "
        "Builds AI-assisted content workflows, immersive visual systems, and social-first creative assets that connect "
        "technical prototyping with campaign execution. Experienced with ComfyUI, Stable Diffusion, Unreal Engine, "
        "Maya, After Effects, DaVinci Resolve, Airtable, and web technologies.",
        size=8.9,
        color=BODY,
    )

    add_section_heading(doc, "Core Skills")
    add_skill_line(
        doc,
        "AI creative and prototyping",
        "ComfyUI, Stable Diffusion, Midjourney, custom LoRA models, AI-assisted previz, face replacement, generative content workflows",
    )
    add_skill_line(
        doc,
        "Creative technology and VFX",
        "Unreal Engine, Blender, Maya, Houdini, Nuke, VFX, real-time rendering, motion graphics, color grading",
    )
    add_skill_line(
        doc,
        "Content and social strategy",
        "Trailer/promo/poster development, short-form video, TikTok and WeChat content, Google Analytics, SEO, KPI analysis, Airtable asset tracking",
    )
    add_skill_line(
        doc,
        "Web and interactive",
        "HTML, CSS, JavaScript, responsive layouts, interactive prototyping, dynamic content, SQL",
    )

    add_section_heading(doc, "Experience")
    add_role(
        doc,
        "Highland Film Group (The Avenue)",
        "Los Angeles, CA",
        "US Domestic Marketing and Distribution Intern",
        "Jan 2026 - Present",
    )
    add_bullet(doc, ["Develop creative assets for theatrical and digital releases, including trailers, promos, posters, and campaign visual materials."])
    add_bullet(doc, ["Translate campaign concepts into executable visual deliverables across creative, marketing, and distribution workflows."])
    add_bullet(doc, ["Manage marketing and creative assets in Airtable, supporting version control, asset accounting, and cross-team visibility across film titles."])
    add_bullet(doc, ["Built an AI agent to support social media content creation, improving marketing-team workflow efficiency by ", ("at least 30%", True), "."])

    add_role(doc, "F.O.U.N.D.", "New York, NY", "Web Design Artist (Part-time)", "Sep 2025 - Present")
    add_bullet(doc, ["Designed responsive web layouts and interactive visual components with modular systems, scalability, and visual consistency."])
    add_bullet(doc, ["Implemented JavaScript-driven interaction and dynamic content to support richer user experiences and rapid creative iteration."])

    add_role(doc, "JOYME", "Beijing, China", "Technical Artist", "Mar 2025 - Jun 2025")
    add_bullet(doc, ["Designed AI-assisted VFX pipelines for face replacement, character enhancement, and asset reuse, improving production efficiency by ", ("around 30%", True), "."])
    add_bullet(doc, ["Integrated generative AI outputs into live-action footage using ComfyUI and node-based compositing workflows."])
    add_bullet(doc, ["Led research and design teams in applying generative AI to previsualization, concept iteration, and look development."])

    add_role(doc, "NetEase", "Beijing, China", "Video Editor", "Feb 2025 - Mar 2025")
    add_bullet(doc, ["Produced short-form cinematic content with VFX, motion graphics, and color grading for fast-turnaround social platforms."])
    add_bullet(doc, ["Edited in DaVinci Resolve and After Effects under tight delivery timelines."])
    add_bullet(doc, ["Created video clips that generated ", ("100,000+ likes", True), " across platforms including TikTok and WeChat Official Accounts."])

    add_role(doc, "Studio KAY", "Busan, South Korea", "Graphic Artist Intern", "Jun 2023 - Jan 2024")
    add_bullet(doc, ["Built digital human assets and VFX elements for XR and immersive media campaigns."])
    add_bullet(doc, ["Optimized real-time rendering workflows in Unreal Engine and Maya, reducing render time by ", ("20%", True), "."])
    add_bullet(doc, ["Collaborated with VFX supervisors to align shader, lighting, and rendering setups with narrative intent."])

    add_section_heading(doc, "Selected Project")
    add_role(doc, "RealLife AI", "Personal Project", "AI Photo Enhancement Platform", "ComfyUI and custom LoRA models")
    add_bullet(doc, ["Built a web-based AI photo enhancement platform designed for high-fidelity photorealism and identity preservation."])
    add_bullet(doc, ["Focused the product experience on lowering technical barriers so nontechnical users can access AI-powered visual enhancement."])

    add_section_heading(doc, "Education")
    education = [
        ("Syracuse University", "Syracuse, NY, USA", "M.S. in Advanced Media Management | Expected Dec 2026"),
        ("Dongseo University", "Busan, South Korea", "B.E. in Film & VFX | GPA: 3.89 / 4.0"),
        ("Zhongnan University of Economics and Law", "Wuhan, China", "B.A. in Film Studies | GPA: 3.7 / 4.0"),
    ]
    for school, place, detail in education:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1.5)
        p.paragraph_format.space_after = Pt(0)
        add_run(p, school, bold=True, size=8.8, color=BODY)
        add_run(p, f" - {place}", size=8.8, color=BODY)
        d = doc.add_paragraph()
        d.paragraph_format.space_before = Pt(0)
        d.paragraph_format.space_after = Pt(1)
        add_run(d, detail, italic=True, size=8.6, color=MUTED)

    doc.core_properties.title = "Jingchuan Carlton Shi - AI Creative and Content Strategy Resume"
    doc.core_properties.subject = "Resume optimized for AI Creative, Creative Technologist, Social Media, and Content Strategy roles"
    doc.core_properties.author = "Jingchuan Carlton Shi"

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
